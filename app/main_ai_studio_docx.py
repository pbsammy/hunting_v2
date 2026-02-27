#!/usr/bin/env python3
"""
CTA DOCX Threat Hunt Report Generator
- Calls Gemini to generate CTA report sections in JSON
- Writes directly into a CTA-styled Word template (DOCX)
- Saves raw model output and parsed sections for troubleshooting
"""
import argparse
import os
import sys
import json
import re
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from jinja2 import Template

# ----------------------------
# Logging / filesystem helpers
# ----------------------------
def log(msg: str) -> None:
    ts = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    sys.stderr.write(f"[{ts}] {msg}\n")
    sys.stderr.flush()

def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

# ----------------------------
# Model call + JSON handling
# ----------------------------
def _extract_json(text: str) -> dict:
    """
    Try to pull a JSON object from:
    - ```json ... ``` fenced blocks
    - the first {...} object in free-form text
    """
    # 1) fenced block ```json ... ```
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=(re.DOTALL | re.IGNORECASE))
    if m:
        return json.loads(m.group(1))
    # 2) first JSON object {...}
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if m:
        return json.loads(m.group(0))
    raise RuntimeError("No JSON object found in model output.")

def call_model(api_key: str, system_prompt: str, user_prompt: str, model: str) -> Dict[str, Any]:
    client = genai.Client(api_key=api_key)

    generation_config = {
        "temperature": 0.2,
        "top_p": 0.9,
        "max_output_tokens": 8192,
        # Strong hint for JSON-only; supported in recent google-genai versions.
        "response_mime_type": "application/json",
        # Optional schema; if unsupported, model often still honors mime type.
        "response_schema": {
            "type": "object",
            "properties": {
                "sections": {
                    "type": "object",
                    "properties": {
                        "background": {"type": "string"},
                        "hypothesis": {"type": "string"},
                        "analysis": {"type": "string"},
                        "findings": {"type": "string"},
                        "recommendations": {"type": "string"},
                        "additional_research": {"type": "string"},
                        "appendix": {"type": "string"},
                        "resources": {"type": "array", "items": {"type": "string"}}
                    },
                    "required": [
                        "background", "hypothesis", "analysis", "findings",
                        "recommendations", "additional_research", "appendix", "resources"
                    ]
                }
            },
            "required": ["sections"]
        }
    }

    # Compose contents (system guidance and the rendered user prompt with IDEA)
    contents = [{
        "role": "user",
        "parts": [
            {"text": f"SYSTEM INSTRUCTION:\n{system_prompt.strip()}"},
            {"text": (
                "Return only a single JSON object with this exact structure:\n"
                "{\n"
                '  "sections": {\n'
                '    "background": "...",\n'
                '    "hypothesis": "...",\n'
                '    "analysis": "...",\n'
                '    "findings": "...",\n'
                '    "recommendations": "...",\n'
                '    "additional_research": "...",\n'
                '    "appendix": "...",\n'
                '    "resources": ["...", "..."]\n'
                "  }\n"
                "}\n"
                "No commentary, no markdown, no code fences, no extra text.\n"
                "THREAT HUNT IDEA:\n"
                f"{user_prompt.strip()}"
            )}
        ]
    }]

    response = client.models.generate_content(
        model=model,
        contents=contents,
        config=generation_config
    )
    if not response or not getattr(response, "text", None):
        raise RuntimeError("Model returned empty response")

    raw = response.text
    ensure_dir("output")
    with open("output/model_raw.txt", "w", encoding="utf-8") as f:
        f.write(raw)

    # Strict parse first; if it fails, attempt extraction from fenced blocks or first {...}
    try:
        return json.loads(raw)
    except Exception:
        try:
            return _extract_json(raw)
        except Exception:
            raise RuntimeError("Model did not return valid JSON")

# ----------------------------
# DOCX helpers (CTA styling)
# ----------------------------
def stamp_header_footer(doc: Document):
    # Add CTA header/footer; avoid clearing if template already has content
    section = doc.sections[0]

    # Header
    hp1 = section.header.add_paragraph()
    r1 = hp1.add_run("TRUST IN DISA – MISSION FIRST, PEOPLE ALWAYS")
    r1.bold = True
    hp1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    hp2 = section.header.add_paragraph()
    r2 = hp2.add_run("CUI//FEDCON")
    r2.bold = True
    hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer
    fp1 = section.footer.add_paragraph(
        "Controlled By: Defense Information Systems Agency (DISA) "
        "DEOS Program Management Office (PMO) (DISA SD3)"
    )
    fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fp2 = section.footer.add_paragraph(
        "CUI Category: General Proprietary Business Information\n"
        "Limited Dissemination Control: Federal Employees and Contractors Only (FEDCON)"
    )
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_cover(doc: Document, prepared_by: str):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Threat Hunt Report").bold = True

    meta = doc.add_paragraph()
    meta.add_run("Controlled Unclassified Information (CUI)\n").bold = True

    doc.add_paragraph(f"Prepared by: {prepared_by}")
    doc.add_paragraph("Document Owner: DEOS Program Management Office")
    doc.add_paragraph("OPR: DISA SD3")
    doc.add_paragraph(
        "CUI DESIGNATION INDICATOR: CUI Category: General Proprietary Business Information\n"
        "Dissemination: FEDCON"
    )

    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "DISCLAIMER\n"
        "The contents of this document are not to be construed as an official Defense Information Systems Agency document "
        "unless so designated by other authorized documents. The use of trade names in this document does not constitute "
        "an official endorsement or approval. Do not cite this document for the purpose of advertisement."
    )

def add_section(doc: Document, title: str, body):
    doc.add_paragraph(title, style="Heading 1")
    if isinstance(body, list):
        for item in body:
            doc.add_paragraph(str(item))
    else:
        doc.add_paragraph(str(body) if body else "[Insert content]")

# ----------------------------
# Main
# ----------------------------
def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument("--system-file", required=True)
    ap.add_argument("--template", required=True)     # must be a DOCX file
    ap.add_argument("--prompt", required=True)       # path to user prompt template (md/txt)
    ap.add_argument("--prepared-by", default="Shawn McWhirter")
    ap.add_argument("--model", default="gemini-2.5-flash")
    ap.add_argument("--output", required=True)
    args = ap.parse_args(argv)

    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        log("ERROR: GEMINI_API_KEY not set")
        return 2

    # Read system prompt
    try:
        with open(args.system_file, "r", encoding="utf-8") as f:
            system_prompt = f.read()
    except Exception as e:
        log(f"ERROR reading system prompt file: {e}")
        return 3

    # Read & render the user prompt template with IDEA context
    idea = os.environ.get("IDEA", "").strip()
    try:
        with open(args.prompt, "r", encoding="utf-8") as f:
            user_prompt_tpl = f.read()
    except Exception as e:
        log(f"ERROR reading user prompt file: {e}")
        return 3

    try:
        # Minimal context injection; expand if you want more fields mapped
        rendered_user_prompt = Template(user_prompt_tpl).render(
            THREAT_NAME=idea or "Threat",
            MITRE_ATTACK_ID="TBD",
            THREAT_DESCRIPTION="",
            ATTACK_VECTOR="",
            DETECTION_HYPOTHESIS="",
            RESOURCES=""
        )
    except Exception as e:
        log(f"ERROR rendering user prompt template: {e}")
        return 3

    # Call LLM for structured sections
    try:
        data = call_model(api_key, system_prompt, rendered_user_prompt, args.model)
    except Exception as e:
        log(str(e))
        return 4

    sections = data.get("sections", {})
    required = [
        "background", "hypothesis", "analysis", "findings",
        "recommendations", "additional_research", "appendix", "resources"
    ]
    missing = [k for k in required if k not in sections]
    if missing:
        log(f"ERROR: missing required sections: {missing}")
        ensure_dir("output")
        with open("output/sections.json", "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        return 5

    # Save structured JSON for inspection
    ensure_dir("output")
    with open("output/sections.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)

    # Load CTA DOCX template and stamp banners/styles
    try:
        doc = Document(args.template)   # expects a .docx file
    except Exception as e:
        log(f"ERROR opening CTA template: {e}")
        return 6

    set_styles(doc)
    stamp_header_footer(doc)
    add_cover(doc, args.prepared_by)

    # Write CTA sections in canonical order
    order = [
        ("Background", sections.get("background")),
        ("Hypothesis", sections.get("hypothesis")),
        ("Analysis", sections.get("analysis")),
        ("Findings", sections.get("findings")),
        ("Recommendations", sections.get("recommendations")),
        ("Additional Research", sections.get("additional_research")),
        ("Appendix", sections.get("appendix")),
        ("Resources", sections.get("resources", [])),
    ]
    for title, body in order:
        add_section(doc, title, body)

    # Save output DOCX
    try:
        doc.save(args.output)
    except Exception as e:
        log(f"ERROR writing DOCX: {e}")
        return 7

    log(f"SUCCESS: wrote CTA DOCX to {args.output}")
    print(json.dumps({"status": "ok", "docx": args.output}, indent=2))
    return 0

if __name__ == "__main__":
    sys.exit(main())
